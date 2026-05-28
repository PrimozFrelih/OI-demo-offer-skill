#!/usr/bin/env python3
"""Convert a Markdown offer into a simple DOCX document."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def set_document_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size in (
        ("Title", 18),
        ("Heading 1", 15),
        ("Heading 2", 13),
        ("Heading 3", 11.5),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text.strip()


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and is_table_row(lines[index]):
        cells = [clean_inline(cell.strip()) for cell in lines[index].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in cells):
            rows.append(cells)
        index += 1
    return rows, index


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=0, cols=column_count)
    table.style = "Table Grid"

    for row_index, row_values in enumerate(rows):
        cells = table.add_row().cells
        for cell_index in range(column_count):
            value = row_values[cell_index] if cell_index < len(row_values) else ""
            paragraph = cells[cell_index].paragraphs[0]
            paragraph.text = value
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
            if re.search(r"(EUR|%|^\d)", value):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_paragraph()


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(clean_inline(text))
    paragraph.paragraph_format.space_after = Pt(6)


def markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    set_document_styles(document)

    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        if is_table_row(stripped):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(clean_inline(stripped[2:]), style="Title")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("## "):
            document.add_heading(clean_inline(stripped[3:]), level=1)
        elif stripped.startswith("### "):
            document.add_heading(clean_inline(stripped[4:]), level=2)
        elif stripped.startswith("#### "):
            document.add_heading(clean_inline(stripped[5:]), level=3)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(2)
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(clean_inline(re.sub(r"^\d+\. ", "", stripped)), style="List Number")
            paragraph.paragraph_format.space_after = Pt(2)
        else:
            add_paragraph(document, stripped)

        index += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown offer to DOCX.")
    parser.add_argument("markdown_file", help="Input Markdown file.")
    parser.add_argument("docx_file", nargs="?", help="Output DOCX file.")
    args = parser.parse_args()

    markdown_path = Path(args.markdown_file)
    if args.docx_file:
        output_path = Path(args.docx_file)
    else:
        output_path = markdown_path.with_suffix(".docx")

    markdown_to_docx(markdown_path, output_path)
    print(f"Built {output_path}")


if __name__ == "__main__":
    main()
